#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
将Markdown文档转换为Word格式
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re
    print('✅ 所有依赖库已导入')
except ImportError as e:
    print(f'❌ 缺少依赖库: {e}')
    print('请运行: python -m pip install python-docx')
    import sys
    sys.exit(1)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    # 设置中文字体
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_paragraph(doc, text, style=None):
    """添加段落"""
    para = doc.add_paragraph(text, style=style)
    # 设置中文字体
    for run in para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return para

def add_code_block(doc, code_text, language=''):
    """添加代码块"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    
    # 设置代码样式
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 设置背景色（浅灰色）
    shading_elm = para._element.get_or_add_pPr()
    shading = shading_elm.get_or_add_shd()
    shading.set(qn('w:fill'), 'F5F5F5')
    
    return para

def add_table_from_markdown(doc, lines):
    """从Markdown表格创建Word表格"""
    # 解析表格行
    rows = []
    for line in lines:
        if '|' in line and not line.strip().startswith('|---'):
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
    
    if not rows:
        return None
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # 填充数据
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_data
                # 设置单元格字体
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                # 表头加粗
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    return table

def convert_markdown_to_word(md_file, docx_file):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block = []
    code_language = ''
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_block = []
            else:
                in_code_block = False
                add_code_block(doc, '\n'.join(code_block), code_language)
                code_block = []
            i += 1
            continue
        
        if in_code_block:
            code_block.append(line)
            i += 1
            continue
        
        # 处理表格
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            # 检查下一行是否还是表格
            if i < len(lines) and '|' not in lines[i]:
                in_table = False
                add_table_from_markdown(doc, table_lines)
                table_lines = []
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 6:
                add_heading(doc, text, level=min(level, 4))
            i += 1
            continue
        
        # 处理分隔线
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue
        
        # 处理普通段落
        if line.strip():
            # 处理粗体和代码
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 移除粗体标记
            text = re.sub(r'`(.*?)`', r'\1', text)  # 移除行内代码标记
            
            para = add_paragraph(doc, text)
            
            # 处理列表项
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                para.style = 'List Bullet'
            elif re.match(r'^\d+\.\s', line.strip()):
                para.style = 'List Number'
        else:
            # 空行
            doc.add_paragraph()
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f'✅ 转换成功！Word文档已保存到: {docx_file}')

if __name__ == '__main__':
    md_file = 'e:/tudianershatest/第3章_WEB开发技术点应用情况.md'
    docx_file = 'e:/tudianershatest/第3章_WEB开发技术点应用情况.docx'
    
    try:
        convert_markdown_to_word(md_file, docx_file)
    except Exception as e:
        print(f'❌ 转换失败: {str(e)}')
        import traceback
        traceback.print_exc()
