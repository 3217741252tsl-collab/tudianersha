# coding: utf-8
import sys
import os

print("=" * 60)
print("Markdown to Word Converter")
print("=" * 60)

# 检查并安装依赖
print("\n[1/4] 检查依赖库...")
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    print("✓ python-docx 已安装")
except ImportError as e:
    print("✗ python-docx 未安装，正在安装...")
    os.system("python -m pip install python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple")
    print("✓ python-docx 安装完成，请重新运行此脚本")
    sys.exit(0)

import re

# 文件路径
md_file = r"e:\tudianershatest\第3章_WEB开发技术点应用情况.md"
docx_file = r"e:\tudianershatest\第3章_WEB开发技术点应用情况.docx"

print(f"\n[2/4] 读取Markdown文件: {md_file}")
try:
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')
    print(f"✓ 成功读取 {len(lines)} 行内容")
except Exception as e:
    print(f"✗ 读取失败: {e}")
    sys.exit(1)

print("\n[3/4] 转换为Word文档...")
try:
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    i = 0
    in_code = False
    code_lines = []
    processed = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                # 结束代码块
                in_code = False
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'No Spacing'
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
            i += 1
            continue
            
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        
        # 标题
        if line.startswith('#'):
            level = 0
            for char in line:
                if char == '#':
                    level += 1
                else:
                    break
            text = line[level:].strip()
            if text:
                h = doc.add_heading(text, level=min(level, 3))
                for run in h.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                processed += 1
        # 普通段落
        elif line.strip():
            # 清理Markdown标记
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text)
            processed += 1
        
        i += 1
        
        # 进度提示
        if i % 500 == 0:
            print(f"  处理进度: {i}/{len(lines)} 行 ({i*100//len(lines)}%)")
    
    print(f"✓ 已处理 {processed} 个段落")
    
except Exception as e:
    print(f"✗ 转换失败: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)

print(f"\n[4/4] 保存Word文档: {docx_file}")
try:
    doc.save(docx_file)
    print(f"✓ 文件大小: {os.path.getsize(docx_file) / 1024:.2f} KB")
except Exception as e:
    print(f"✗ 保存失败: {e}")
    sys.exit(1)

print("\n" + "=" * 60)
print("✅ 转换成功完成！")
print("=" * 60)
print(f"\n输出文件: {docx_file}")
